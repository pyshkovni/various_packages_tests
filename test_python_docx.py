# Эксперименты с библиотекой python-docx

# C чего мы начинаем изучение новой библиотеки?
# 1. Погуглить/поджипидишить ->  генерация, гайд, страница библиотеки, pypi
# 2. Поиск через pypi.org -> страница библиотеки
# 3. Почитать документацию
# https://python-docx.readthedocs.io/en/latest/

# импорты
from docx import Document

def python_docx_tests(path):
    # создание объекта документа
    document = Document()

    print(document) # <docx.document.Document object at 0x7f8f42393fb0>

    # код про добавление текста в документ
    paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')
    prior_paragraph = paragraph.insert_paragraph_before('Lorem ipsum')
    document.add_heading('The REAL meaning of the universe')
    document.add_heading('The role of dolphins', level=2)

    # сохранить в новый документ
    document.save('demo.docx')

    # свойства документа
    document = Document(path)
    properties = document.core_properties
    print("Автор:", properties.author)
    print("Время создания:", properties.created)

